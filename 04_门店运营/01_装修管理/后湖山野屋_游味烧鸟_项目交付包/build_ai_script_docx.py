from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "2026_AI改造烧鸟店_系列短视频脚本.md"
OUT = ROOT / "2026_烧鸟主理人用AI重新改造店铺_系列短视频脚本.docx"
ACCENT = "A0442C"
ACCENT_DARK = "612A20"
INK = "292622"
MUTED = "746D65"
PALE = "F3ECE5"
GOLD = "B88645"


def set_font(run, size=None, bold=None, color=INK, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def border_bottom(paragraph, color=ACCENT, size=14):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("第 ")
    set_font(r, 8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    r = paragraph.add_run(" 页")
    set_font(r, 8.5, color=MUTED)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.88)
    sec.right_margin = Inches(0.88)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 17, ACCENT_DARK, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 11.5, ACCENT_DARK, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("游味烧鸟  ·  AI 改店实录")
    set_font(run, 8.5, True, MUTED)
    border_bottom(header, "D8CCC1", 5)
    add_page_field(sec.footer.paragraphs[0])
    return doc


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026 · 真实改店连续剧")
    set_font(r, 10.5, True, GOLD)
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("烧鸟主理人用 AI\n重新改造一家湖边小店")
    set_font(r, 27, True, ACCENT_DARK)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《游味烧鸟》11 集短视频完整拍摄脚本")
    set_font(r, 13, False, MUTED)
    p.paragraph_format.space_after = Pt(30)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(5.3)
    shade_cell(cell, PALE)
    set_cell_margins(cell, 220, 260, 220, 260)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("不是让 AI 一键开店，\n而是公开验证：它能不能帮一个普通主理人少走弯路。")
    set_font(r, 12, True, ACCENT)
    p.paragraph_format.line_spacing = 1.45

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("湖畔 × 木构 × 炭火 × 夜晚 × 长沙")
    set_font(r, 10.5, True, ACCENT)
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("游于后湖，味在人间")
    set_font(r, 11, False, MUTED)
    doc.add_page_break()


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    r = p.add_run(label + "  ")
    set_font(r, 9.5, True, "FFFFFF")
    r._element.get_or_add_rPr().append(OxmlElement("w:shd"))
    r._element.rPr[-1].set(qn("w:fill"), ACCENT)
    r = p.add_run(text)
    set_font(r, 10.5, False, INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, 10.5)


def clean_md(s):
    return re.sub(r"\*\*(.*?)\*\*", r"\1", s).replace("`", "")


def add_content(doc):
    lines = SRC.read_text(encoding="utf-8").splitlines()
    in_series_intro = True
    first_episode = True
    i = 1
    while i < len(lines):
        raw = lines[i].rstrip()
        s = raw.strip()
        if not s or s == "---":
            i += 1
            continue
        if s.startswith("## 第 "):
            if first_episode:
                doc.add_page_break()
                first_episode = False
            else:
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(clean_md(s[3:]))
            set_font(r, 17, True, ACCENT_DARK)
            border_bottom(p, ACCENT, 10)
            in_series_intro = False
        elif s == "## 系列定位":
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run("系列总览")
            set_font(r, 17, True, ACCENT_DARK)
            border_bottom(p, ACCENT, 10)
        elif s.startswith("## "):
            doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(clean_md(s[3:]))
            set_font(r, 17, True, ACCENT_DARK)
            border_bottom(p, ACCENT, 10)
        elif s.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(clean_md(s[4:]))
            set_font(r, 13, True, ACCENT)
        elif s.startswith("**建议时长："):
            add_label_para(doc, "建议时长", clean_md(s).replace("建议时长：", ""))
        elif s.startswith("**封面标题**："):
            add_label_para(doc, "封面标题", clean_md(s).replace("封面标题：", ""))
        elif s.startswith("**开场 ") or s.startswith("**正文 ") or s.startswith("**结尾 ") or s.startswith("**收束"):
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(clean_md(s))
            set_font(r, 13, True, ACCENT)
        elif s.startswith("【画面】"):
            add_label_para(doc, "画面", clean_md(s.replace("【画面】", "")))
        elif s.startswith("【口播】"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("口播")
            set_font(r, 10, True, GOLD)
        elif s.startswith("- "):
            add_bullet(doc, clean_md(s[2:]))
        elif re.match(r"^\d+\. ", s):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(clean_md(re.sub(r"^\d+\. ", "", s)))
            set_font(r, 10.5)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.widow_control = True
            text = clean_md(s)
            if text.startswith("“"):
                p.paragraph_format.left_indent = Inches(0.16)
                p.paragraph_format.right_indent = Inches(0.12)
                r = p.add_run(text)
                set_font(r, 10.7, False, INK)
            else:
                r = p.add_run(text)
                set_font(r, 10.5)
        i += 1


def main():
    doc = setup_doc()
    add_cover(doc)
    add_content(doc)
    props = doc.core_properties
    props.title = "2026 烧鸟主理人用 AI 重新改造店铺｜系列短视频脚本"
    props.subject = "游味烧鸟真实改店连续剧完整拍摄手册"
    props.author = "游味烧鸟"
    props.keywords = "AI开店, 烧鸟, 主理人, 短视频脚本, 游味烧鸟"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
