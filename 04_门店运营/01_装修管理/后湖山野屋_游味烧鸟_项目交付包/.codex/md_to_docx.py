from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(sys.argv[1]).resolve()

LATIN_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
HEADING_BLUE = "2E74B5"
HEADING_DARK_BLUE = "1F4D78"
MUTED = "6B7280"


def set_run_font(run, size=None, bold=None, italic=None, color=None, mono=False):
    latin = "Consolas" if mono else LATIN_FONT
    east = "Microsoft YaHei" if not mono else "DengXian"
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key, value in (("ascii", latin), ("hAnsi", latin), ("eastAsia", east), ("cs", latin)):
        rfonts.set(qn(f"w:{key}"), value)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    heading_tokens = {
        1: (16, HEADING_BLUE, 18, 10),
        2: (13, HEADING_BLUE, 14, 7),
        3: (12, HEADING_DARK_BLUE, 10, 5),
        4: (11.5, HEADING_DARK_BLUE, 8, 4),
        5: (11, HEADING_DARK_BLUE, 7, 3),
        6: (10.5, HEADING_DARK_BLUE, 6, 3),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else f"%{level + 1}.")
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abs_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def set_list_numbering(paragraph, num_id: int, level: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(numid)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


INLINE_RE = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|(?<!\*)\*[^*]+\*(?!\*)|(?<!_)_[^_]+_(?!_)|\[[^\]]+\]\([^\)]+\))"
)


def add_inline(paragraph, text: str):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            set_run_font(paragraph.add_run(text[pos:m.start()]))
        token = m.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=10, mono=True, color="374151")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F3F4F6")
            run._element.get_or_add_rPr().append(shd)
        elif token.startswith("**") or token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("*") or token.startswith("_"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, italic=True)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).groups()
            run = paragraph.add_run(label)
            set_run_font(run, color="0563C1")
            run.underline = True
            run._element.set(qn("w:rsidRPr"), "00000000")
            # Keep the destination visible for reliable offline browsing.
            set_run_font(paragraph.add_run(f"（{url}）"), size=9, color=MUTED)
        pos = m.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]))


def add_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    ppr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    borders.append(bottom)
    ppr.append(borders)


def add_footer(doc: Document):
    for section in doc.sections:
        section.footer_distance = Inches(0.492)
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("第 ")
        set_run_font(run, size=9, color=MUTED)
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)
        run = p.add_run(" 页")
        set_run_font(run, size=9, color=MUTED)


def decode_markdown(path: Path) -> str:
    data = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def convert(path: Path) -> Path:
    text = decode_markdown(path).replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    configure_styles(doc)
    bullet_id = add_numbering(doc, "bullet")
    decimal_id = add_numbering(doc, "decimal")

    i = 0
    para_buf: list[str] = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            p = doc.add_paragraph()
            add_inline(p, " ".join(x.strip() for x in para_buf))
            para_buf = []

    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            flush_para()
            i += 1
            continue
        h = re.match(r"^(#{1,6})\s+(.+?)\s*#*\s*$", line)
        ul = re.match(r"^(\s*)[-+*]\s+(.+)$", line)
        ol = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if h:
            flush_para()
            p = doc.add_paragraph(style=f"Heading {len(h.group(1))}")
            add_inline(p, h.group(2))
        elif ul or ol:
            flush_para()
            m = ul or ol
            level = min(len(m.group(1).replace("\t", "    ")) // 2, 2)
            p = doc.add_paragraph()
            add_inline(p, m.group(2))
            set_list_numbering(p, bullet_id if ul else decimal_id, level)
        elif re.match(r"^\s*([-*_])(?:\s*\1){2,}\s*$", line):
            flush_para()
            add_rule(doc)
        elif line.lstrip().startswith(">"):
            flush_para()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.1)
            p.paragraph_format.space_after = Pt(6)
            ppr = p._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "CBD5E1")
            borders.append(left)
            ppr.append(borders)
            add_inline(p, line.lstrip()[1:].strip())
            for run in p.runs:
                run.font.color.rgb = RGBColor.from_string("4B5563")
        else:
            para_buf.append(line)
        i += 1
    flush_para()
    add_footer(doc)

    target = path.with_suffix(".docx")
    doc.save(target)
    return target


def main():
    markdown_files = sorted(p for p in ROOT.rglob("*.md") if ".codex" not in p.parts)
    outputs = []
    for source in markdown_files:
        target = convert(source)
        outputs.append((source.relative_to(ROOT), target.relative_to(ROOT), target.stat().st_size))
    print(f"converted={len(outputs)}")
    for source, target, size in outputs:
        print(f"{source}\t{target}\t{size}")


if __name__ == "__main__":
    main()
