from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
V2 = ROOT / "04_参考图提取版_V2"
OUT = V2 / "游味烧鸟_VI设计总结_V2.docx"

BRAND_BLUE = RGBColor(0x48, 0x61, 0x6B)
BRAND_BROWN = RGBColor(0x5D, 0x44, 0x2F)
BRAND_ORANGE = RGBColor(0xDD, 0x82, 0x45)
INK = RGBColor(0x22, 0x22, 0x20)
MUTED = RGBColor(0x97, 0x82, 0x6E)
PAPER = "F5F0E8"


def set_run_font(run, size=None, bold=None, color=None, latin="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=8, color=MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, before, after, color in (
    ("Heading 1", 16, 18, 10, BRAND_BLUE),
    ("Heading 2", 13, 14, 7, BRAND_BLUE),
    ("Heading 3", 12, 10, 5, BRAND_BROWN),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = header.add_run("游味烧鸟 VI 视觉系统  |  当前采用 V2")
set_run_font(r, size=8.5, color=MUTED)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run("YOU WEI YAKITORI   ·   ")
set_run_font(r, size=8, color=MUTED)
add_page_field(footer)

# Cover: editorial_cover pattern with brand overrides.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(52)
p.paragraph_format.space_after = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VISUAL IDENTITY SYSTEM")
set_run_font(r, size=10, bold=True, color=BRAND_ORANGE, latin="Montserrat")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("游味烧鸟")
set_run_font(r, size=30, bold=True, color=INK, east_asia="Microsoft YaHei")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("VI 视觉识别系统 · 参考图提取版 V2")
set_run_font(r, size=14, color=BRAND_BLUE)

logo = V2 / "02_可用标志" / "09_标准组合_清稿版_浅底预览.png"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
p.add_run().add_picture(str(logo), width=Inches(5.75))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("后湖有风，炭火有味，")
set_run_font(r, size=14, bold=True, color=BRAND_BROWN)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("在人间烟火里，游味一晚。")
set_run_font(r, size=14, bold=True, color=BRAND_BROWN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(28)
r = p.add_run("CURRENT APPROVED DIRECTION  ·  2026")
set_run_font(r, size=9, color=MUTED, latin="Montserrat")

doc.add_page_break()

doc.add_heading("1. 当前采用方向", level=1)
p = doc.add_paragraph()
r = p.add_run("本版本以用户确认的 VI 参考总图为唯一视觉依据。")
set_run_font(r, bold=True, color=BRAND_BROWN)
p.add_run(" 此前原创“湖火签”方案停止作为执行标准，常用文件已经同步至原有 Logo、应用效果图及辅助规范目录。")

ref = V2 / "01_原图提取" / "02_VI基础系统总览.png"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
p.add_run().add_picture(str(ref), width=Inches(6.3))

doc.add_heading("2. 标志系统", level=1)
doc.add_heading("中文标准字", level=2)
doc.add_paragraph("采用参考图中的手写书法字形“游味烧鸟”。字形松弛、有手作温度，与后湖木屋、炭火和夜间灯光相符。品牌书法字必须使用提取图形，不能用普通字体替代。")
doc.add_heading("圆形品牌章", level=2)
doc.add_paragraph("由书法标准字、圆形轮廓、湖面波纹与红色印章组成。适用于灯笼、杯具、外带袋、围裙、菜单扉页、封签和社交头像。")
doc.add_heading("标准组合", level=2)
doc.add_paragraph("中文标准字为第一层级；YOU WEI YAKITORI 为第二层级；品类描述统一使用“后湖湖畔 · 炭火烧鸟”。")

doc.add_page_break()

doc.add_heading("3. 色彩与字体", level=1)
palette = V2 / "03_规范整理" / "01_品牌色彩_提取版.png"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(palette), width=Inches(6.25))

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Inches(1.35), Inches(1.3), Inches(3.85)]
headers = ["色名", "HEX", "主要用途"]
for i, (cell, text, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
    cell.width = width
    cell.text = text
    shade_cell(cell, "48616B")
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for run in cell.paragraphs[0].runs:
        set_run_font(run, size=9.5, bold=True, color=RGBColor(255, 255, 255))
set_repeat_table_header(table.rows[0])
rows = [
    ("湖水青", "#48616B", "湖景传播、杯套、信息背景"),
    ("暖肌纸", "#CDC3BA", "菜单、包装与纸张底色"),
    ("竹苔绿", "#606042", "植物、户外及季节物料"),
    ("炭火黑", "#222220", "标志、正文、制服与门头"),
    ("木原褐", "#5D442F", "原木、围裙、外带包装"),
    ("暖阳橙", "#DD8245", "印章、炭火和重点提示"),
    ("茶白", "#D7CABC", "内页与深色背景反白"),
    ("城市棕", "#97826E", "次级文字、边框与材质信息"),
]
for idx, row_data in enumerate(rows):
    cells = table.add_row().cells
    for i, (cell, text, width) in enumerate(zip(cells, row_data, widths)):
        cell.width = width
        cell.text = text
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if idx % 2 == 1:
            shade_cell(cell, "F1ECE4")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.2, color=INK)

doc.add_heading("字体规范", level=2)
doc.add_paragraph("中文标题：汉仪手书体 / 粗；中文正文：思源宋体 / 常规；英文：Montserrat / Regular；数字：DIN / Regular。")

doc.add_page_break()

doc.add_heading("4. 辅助图形与材质", level=1)
aux = V2 / "03_规范整理" / "02_辅助图形_提取重绘.png"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(aux), width=Inches(6.2))
doc.add_paragraph("湖纹、远山与两笔式飞鸟只作轻量点缀；红印作为视觉落款，每个画面原则上只出现一次。")

materials = V2 / "01_原图提取" / "08_材质规范.png"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(materials), width=Inches(5.7))
doc.add_paragraph("材质统一使用原木、石材、炭化木、绿植与麻布，强调自然、克制、松弛，避免堆砌传统日式符号。")

doc.add_heading("5. 应用系统", level=1)
applications = V2 / "01_原图提取" / "10_周边与品牌用品.png"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(applications), width=Inches(6.3))
doc.add_paragraph("应用覆盖名片、纸杯、打包袋、烧鸟签、围裙、员工服、灯笼、菜单夹、桌牌、酒瓶标签、外卖盒及礼品盒。")

doc.add_heading("6. 生产提示", level=1)
p = doc.add_paragraph()
r = p.add_run("重要：")
set_run_font(r, bold=True, color=BRAND_ORANGE)
p.add_run("原始参考图是 1216×1294 位图。当前透明底标志可用于方案确认、菜单、杯具、包装和常规屏幕；大型门头、灯箱或高精度印刷前，需要依据当前字形进行人工矢量描摹，并按材质实物校色。")

doc.save(OUT)
print(OUT)
